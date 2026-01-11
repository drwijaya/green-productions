"""Service for exporting DSO to Word document."""
import os
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from .barcode_service import generate_qr_code


def export_dso_to_word(dso, use_libre_template=False):
    """Export DSO data to Word document using template.
    
    Args:
        dso: The DSO object to export
        use_libre_template: If True, use LibreOffice-optimized template (for PDF export)
                           If False, use MS Word template (for .docx download)
    """
    if use_libre_template:
        template_name = 'dsotemplates_libre.docx'
    else:
        template_name = 'dsotemplates.docx'
    
    template_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 
                                  '..', 'template_docs', template_name)
    
    if not os.path.exists(template_path):
        # Fallback to original template if libre version doesn't exist
        template_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 
                                      '..', 'template_docs', 'dsotemplates.docx')
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    
    # Load template
    doc = Document(template_path)
    
    # Prepare data mappings
    order = dso.order
    dewasa = dso.size_chart_dewasa
    anak = dso.size_chart_anak
    
    # Build replacement dictionary
    replacements = {
        # Header Info
        '{{INV}}': order.order_code or '',
        '{{DUE DATE}}': order.deadline.strftime('%d/%m/%Y') if order.deadline else '',
        
        # Product Info
        '{{MODEL}}': order.model or '',
        '{{JENIS}}': dso.jenis or '',
        '{{BAHAN}}': dso.bahan or '',
        '{{SABLON}}': dso.sablon or '',
        '{{WARNA}}': dso.warna or '',
        '{{POSISI}}': dso.posisi or '',
        
        # Accessories
        '{{ACC1}}': dso.acc_1 or '',
        '{{ACC2}}': dso.acc_2 or '',
        '{{ACC3}}': dso.acc_3 or '',
        '{{ACC4}}': dso.acc_4 or '',
        '{{ACC5}}': dso.acc_5 or '',
        '{{KANCING}}': dso.kancing or '',
        '{{SAKU}}': dso.saku or '',
        '{{RESLETING}}': dso.resleting or '',
        '{{MODEL BADAN BAWAH}}': dso.model_badan_bawah or '',
        
        # Catatan Customer
        '{{CATATAN CUSTOMER 1}}': dso.catatan_customer_1 or '',
        '{{CATATAN CUSTOMER 2}}': dso.catatan_customer_2 or '',
        '{{CATATAN CUSTOMER 3}}': dso.catatan_customer_3 or '',
        '{{CATATAN CUSTOMER 4}}': dso.catatan_customer_4 or '',
        '{{CATATAN CUSTOMER 5}}': dso.catatan_customer_5 or '',
        '{{CATATAN CUSTOMER 6}}': dso.catatan_customer_6 or '',
        '{{LABEL}}': dso.label or '',
        
        # Customer Info (for QC Sheet page)
        '{{CUSTOMER}}': order.customer.name if order.customer else '',
        
        # Size Chart Dewasa - Pendek
        '{{AD}}': str(dewasa.pendek_xs if dewasa else 0),
        '{{BD}}': str(dewasa.pendek_s if dewasa else 0),
        '{{CD}}': str(dewasa.pendek_m if dewasa else 0),
        '{{DD}}': str(dewasa.pendek_l if dewasa else 0),
        '{{ED}}': str(dewasa.pendek_xl if dewasa else 0),
        '{{FD}}': str(dewasa.pendek_xxl if dewasa else 0),
        '{{GD}}': str(dewasa.pendek_x3l if dewasa else 0),
        '{{HD_PENDEK}}': str(dewasa.pendek_x4l if dewasa else 0),
        '{{ID_PENDEK}}': str(dewasa.pendek_x5l if dewasa else 0),
        
        # Size Chart Dewasa - Panjang
        '{{HD}}': str(dewasa.panjang_xs if dewasa else 0),
        '{{ID}}': str(dewasa.panjang_s if dewasa else 0),
        '{{JD}}': str(dewasa.panjang_m if dewasa else 0),
        '{{KD}}': str(dewasa.panjang_l if dewasa else 0),
        '{{LD}}': str(dewasa.panjang_xl if dewasa else 0),
        '{{MD}}': str(dewasa.panjang_xxl if dewasa else 0),
        '{{ND}}': str(dewasa.panjang_x3l if dewasa else 0),
        '{{OD}}': str(dewasa.panjang_x4l if dewasa else 0),
        
        # Dewasa totals
        '{{JUMDA}}': str(dewasa.jum_pendek if dewasa else 0),
        '{{JUMDB}}': str(dewasa.jum_panjang if dewasa else 0),
        '{{QTYD }}': str(dewasa.total if dewasa else 0),
        '{{QTYD}}': str(dewasa.total if dewasa else 0),
        
        # Size Chart Anak - Pendek
        '{{AA}}': str(anak.pendek_xs if anak else 0),
        '{{BA}}': str(anak.pendek_s if anak else 0),
        '{{CA}}': str(anak.pendek_m if anak else 0),
        '{{DA}}': str(anak.pendek_l if anak else 0),
        '{{EA}}': str(anak.pendek_xl if anak else 0),
        '{{FA}}': str(anak.pendek_xxl if anak else 0),
        '{{GA}}': str(anak.pendek_x3l if anak else 0),
        '{{HA_PENDEK}}': str(anak.pendek_x4l if anak else 0),
        '{{IA_PENDEK}}': str(anak.pendek_x5l if anak else 0),
        
        # Size Chart Anak - Panjang
        '{{HA}}': str(anak.panjang_xs if anak else 0),
        '{{IA}}': str(anak.panjang_s if anak else 0),
        '{{JA}}': str(anak.panjang_m if anak else 0),
        '{{KA}}': str(anak.panjang_l if anak else 0),
        '{{LA}}': str(anak.panjang_xl if anak else 0),
        '{{MA}}': str(anak.panjang_xxl if anak else 0),
        '{{NA}}': str(anak.panjang_x3l if anak else 0),
        '{{OA}}': str(anak.panjang_x4l if anak else 0),
        
        # Anak totals
        '{{JUMAA}}': str(anak.jum_pendek if anak else 0),
        '{{JUMAB}}': str(anak.jum_panjang if anak else 0),
        '{{QTYA }}': str(anak.total if anak else 0),
        '{{QTYA}}': str(anak.total if anak else 0),
    }
    
    # Replace text in all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        # Replace while preserving formatting
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, value)
                            # Also check full paragraph text if runs don't have it
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)
    
    # Replace in paragraphs outside tables
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
    
    # Handle image placeholder {{GAMBAR}}
    if dso.gambar_depan_url:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '{{GAMBAR}}' in cell.text:
                        # Clear the cell
                        for paragraph in cell.paragraphs:
                            paragraph.clear()
                        
                        # Try to add image
                        try:
                            # Download image from URL
                            response = requests.get(dso.gambar_depan_url, timeout=10)
                            if response.status_code == 200:
                                image_stream = io.BytesIO(response.content)
                                # Add image to cell
                                paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                                run = paragraph.add_run()
                                run.add_picture(image_stream, width=Inches(4))
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            print(f"Error adding image: {e}")
                            cell.paragraphs[0].add_run("(Gambar tidak tersedia)")
    else:
        # Remove placeholder if no image
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '{{GAMBAR}}' in cell.text:
                        for paragraph in cell.paragraphs:
                            paragraph.text = paragraph.text.replace('{{GAMBAR}}', '')
    
    # Handle QR Invoice placeholder {{QRINV}}
    if order.order_code:
        qr_buffer = generate_qr_code(order.order_code, size=6)
        if qr_buffer:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if '{{QRINV}}' in cell.text:
                            # Clear the cell
                            for paragraph in cell.paragraphs:
                                paragraph.clear()
                            
                            # Add QR code image
                            try:
                                paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                                run = paragraph.add_run()
                                run.add_picture(qr_buffer, width=Inches(0.65))
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception as e:
                                print(f"Error adding QR code: {e}")
                                cell.paragraphs[0].add_run("(QR tidak tersedia)")
    else:
        # Remove placeholder if no order code
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '{{QRINV}}' in cell.text:
                        for paragraph in cell.paragraphs:
                            paragraph.text = paragraph.text.replace('{{QRINV}}', '')
    
    # Save to BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def export_dso_to_pdf(dso):
    """Export DSO data to PDF document.
    
    On Windows: uses docx2pdf for Word-to-PDF conversion (best fidelity).
    On Linux: uses WeasyPrint to generate PDF from HTML (Railway compatible).
    """
    import sys
    import platform
    
    # Check if we're on Windows
    is_windows = platform.system() == 'Windows'
    
    if is_windows:
        return _export_dso_to_pdf_windows(dso)
    else:
        return _export_dso_to_pdf_weasyprint(dso)


def _export_dso_to_pdf_windows(dso):
    """Export DSO to PDF using docx2pdf (Windows only)."""
    import tempfile
    import site
    import os

    # Fix for pywin32 DLL load error (pywintypes314.dll not found)
    try:
        import sys
        paths_to_check = sys.path + site.getsitepackages()
        for path in paths_to_check:
            dll_path = os.path.join(path, 'pywin32_system32')
            if os.path.exists(dll_path):
                os.environ['PATH'] = dll_path + os.pathsep + os.environ['PATH']
                break
    except Exception as e:
        print(f"Warning: Could not patch PATH for pywin32: {e}")

    try:
        from docx2pdf import convert
        import pythoncom
    except ImportError as e:
        raise ImportError(f"PDF generation requires docx2pdf and pywin32. Error: {e}")
    
    word_buffer = export_dso_to_word(dso)
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
        tmp_docx.write(word_buffer.getvalue())
        tmp_docx_path = tmp_docx.name
    
    tmp_pdf_path = tmp_docx_path.replace('.docx', '.pdf')
    
    try:
        pythoncom.CoInitialize()
        convert(tmp_docx_path, tmp_pdf_path)
        
        with open(tmp_pdf_path, 'rb') as f:
            pdf_content = f.read()
            
        pdf_buffer = io.BytesIO(pdf_content)
        return pdf_buffer
        
    finally:
        if os.path.exists(tmp_docx_path):
            os.remove(tmp_docx_path)
        if os.path.exists(tmp_pdf_path):
            os.remove(tmp_pdf_path)


def _export_dso_to_pdf_weasyprint(dso):
    """Export DSO to PDF using LibreOffice (Linux compatible).
    
    Uses the LibreOffice-optimized template for better PDF output.
    """
    import tempfile
    import subprocess
    import os
    
    # Generate the Word document using the LibreOffice-optimized template
    word_buffer = export_dso_to_word(dso, use_libre_template=True)
    
    # Create temp directory for conversion
    with tempfile.TemporaryDirectory() as tmpdir:
        # Save Word file
        docx_path = os.path.join(tmpdir, 'dso.docx')
        pdf_path = os.path.join(tmpdir, 'dso.pdf')
        
        with open(docx_path, 'wb') as f:
            f.write(word_buffer.getvalue())
        
        # Convert using LibreOffice headless
        lo_commands = ['libreoffice', 'soffice']
        converted = False
        
        for lo_cmd in lo_commands:
            try:
                result = subprocess.run([
                    lo_cmd,
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', tmpdir,
                    docx_path
                ], capture_output=True, text=True, timeout=60)
                
                if result.returncode == 0 and os.path.exists(pdf_path):
                    converted = True
                    break
            except FileNotFoundError:
                continue
        
        if not converted:
            raise Exception("PDF conversion failed. Please install LibreOffice: sudo pacman -S libreoffice-fresh")
        
        # Read the generated PDF
        with open(pdf_path, 'rb') as f:
            pdf_content = f.read()
    
    pdf_buffer = io.BytesIO(pdf_content)
    return pdf_buffer

